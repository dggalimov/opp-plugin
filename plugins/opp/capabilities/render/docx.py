"""cap-render / docx — канонический md (см. engine.py) → docx-представление для приёмки заказчиком.

Тема оформления берётся ТОЛЬКО из `reference/theme.py: active_palette` (контракт reg-themes) —
никаких захардкоженных цветов; сменить тему проекта (`Оформление.yaml`) меняет вид документа
без правки кода. Формат страницы — A4 ландшафт (широкие таблицы протоколов). Колонтитул —
«проект · документ · стр. N»; дата генерации документа пишется ТОЛЬКО в колонтитул (в md-скелет,
проверяемый на байт-стабильность, дата не входит — иначе рендер был бы недетерминирован).

Разбор md ограничен тем, что порождает `engine.py` (заголовки #/##/###, таблицы, абзацы с
**жирным**/списками «- », «---», блок картинки «![подпись](путь)» — кадры демонстраций, спека
06-b) — полноценный markdown-парсер не требуется, симметрии с `interface/md.py` (HTML-предпросмотр
в окне) достаточно для целей приёмки. Путь картинки в md — относительно рабочего пространства
проекта (тот же корень, что и «Документы/»).

Витрина корпоративного документа (спека 05-b-report-asis, редакция 2): титульный лист (название,
опционально клиент/дата — параметры render_docx) + страница «Содержание» с полем TOC (Word
предлагает обновить его при открытии — это штатное поведение самого Word, не дефект рендера) —
ОБА по ключу конфига «документ.титул» (спека 08 §8): рабочие/рассылочные документы (повестка, план
встреч, открытые вопросы, гипотезы) идут с `титул=False` и не несут ни титула, ни «Содержания»;
разрыв страницы перед каждым «## »-разделом. Стили: H2 — крупный текст цвета акцента, H3 — среднего
размера тёмный; таблицы — шапка с заливкой темы + лёгкое чередование фона строк (полосатая заливка),
ширины колонок — по содержимому (см. `_column_widths`).
"""

from __future__ import annotations
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

from reference.theme import active_palette

_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_TABLE_SEP = re.compile(r"^\|[\s:\-|]+\|$")
_IMAGE = re.compile(r"^!\[(.*)\]\((.*)\)$")
_BULLET = re.compile(r"^[-*]\s+(.*)$")
_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip().replace("\\|", "|").replace("\\\\", "\\") for c in s.split("|")]


def _parse_blocks(md_text: str) -> list[dict]:
    """Md (подмножество engine.py) → список блоков: {type: heading|table|text, ...}."""
    lines = md_text.replace("\r\n", "\n").split("\n")
    blocks: list[dict] = []
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue
        m = _HEADING.match(stripped)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2)})
            i += 1
            continue
        if stripped.startswith("_") and stripped.endswith("_"):
            blocks.append({"type": "text", "text": stripped.strip("_"), "italic": True})
            i += 1
            continue
        m_img = _IMAGE.match(stripped)
        if m_img:
            blocks.append({"type": "image", "caption": m_img.group(1), "path": m_img.group(2)})
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < n and _TABLE_SEP.match(lines[i + 1].strip()):
            header = _split_row(lines[i])
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i]))
                i += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue
        if stripped.startswith(">"):
            # блок-вывод (callout, приём отчётов Big4): подряд идущие «> …» строки —
            # выделенная врезка (заливка + акцентная левая граница); «**жирная**» первая
            # строка становится подзаголовком врезки («Ключевой вывод», «Бизнес-эффект»)
            callout: list[str] = []
            while i < n and lines[i].strip().startswith(">"):
                callout.append(lines[i].strip()[1:].strip())
                i += 1
            blocks.append({"type": "callout", "lines": [c for c in callout if c]})
            continue
        m_bullet = _BULLET.match(stripped)
        if m_bullet:
            blocks.append({"type": "bullet", "text": m_bullet.group(1)})
            i += 1
            continue
        blocks.append({"type": "text", "text": stripped})
        i += 1
    return blocks


_BASE_FONT = "Google Sans"   # основной шрифт документа (типографика 06.07); нет в системе → Word подставит замену
_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\s][^*]*?\*)")   # **жирный** и *курсив*


def _font(run, size_pt: float | None = None) -> None:
    """Применить основной шрифт (и размер) к прогону — включая rFonts для полного покрытия."""
    run.font.name = _BASE_FONT
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), _BASE_FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _add_inline_runs(paragraph, text: str, size_pt: float | None = None) -> None:
    """Разбить текст абзаца на прогоны: «**жирный**», «*курсив*» (UX-типографика, 06.07);
    каждому прогону — основной шрифт документа (и размер, если задан)."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            _font(paragraph.add_run(text[pos:m.start()]), size_pt)
        token = m.group(1)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.font.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.italic = True
        _font(run, size_pt)
        pos = m.end()
    if pos < len(text):
        _font(paragraph.add_run(text[pos:]), size_pt)


def _setup_page(doc: Document) -> None:
    """A4 ландшафт — таблицы протоколов широкие (много колонок)."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Mm(297), Mm(210)
    section.left_margin = section.right_margin = Mm(15)
    section.top_margin = section.bottom_margin = Mm(15)


def _add_footer(doc: Document, project: str, document_title: str) -> None:
    """Колонтитул «проект · документ · стр. N»; дата генерации — только здесь (не в md-скелете)."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = date.today().isoformat()
    run = p.add_run(f"{document_title} · {today} · стр. ")
    _font(run, 8)

    # автополе номера страницы (PAGE)
    fld_begin = p._p.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    instr = p._p.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = "PAGE"
    fld_end = p._p.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run2 = p.add_run()
    _font(run2, 8)
    run2._r.append(fld_begin)
    run2._r.append(instr)
    run2._r.append(fld_end)


def _style_heading(paragraph, level: int, palette: dict) -> None:
    """Иерархия корпоративного документа: H2 раздел (акцент, крупно) · H3 функция/паспорт
    (тёмный) · H4 процесс · H5 подпроцесс; # — титул документа (акцент)."""
    paragraph.style = f"Heading {min(level, 9)}" if level > 1 else "Title"
    color = _rgb(palette["text"]) if level >= 3 else _rgb(palette["accent"])
    for run in paragraph.runs:
        run.font.color.rgb = color
        _font(run)
    _sizes = {2: 16, 3: 13, 4: 12, 5: 11}
    _spaces = {2: (18, 8), 3: (12, 6), 4: (10, 4), 5: (8, 4)}
    if level in _sizes:
        before, after = _spaces[level]
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        for run in paragraph.runs:
            run.font.size = Pt(_sizes[level])


def _page_break_before(paragraph) -> None:
    """Разрыв страницы перед разделом «## » — корпоративный документ, один раздел = одна страница-старт."""
    paragraph.paragraph_format.page_break_before = True


_TABLE_FONT_DELTA = 1.0     # размер шрифта таблиц = основной − 1 (типографика 06.07)
_BR = re.compile(r"<br\s*/?>", re.IGNORECASE)


def _fill_cell(cell, text: str, size_pt: float, bold_all: bool = False,
               color=None) -> None:
    """Ячейка таблицы: инлайн-разметка (**жирный**/*курсив*), переносы «<br>» — отдельными
    абзацами, единый шрифт и размер; bold_all/color — для шапки и первой колонки."""
    parts = _BR.split(text) or [""]
    first = True
    for part in parts:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(2)
        _add_inline_runs(p, part.strip(), size_pt)
        if bold_all or color is not None:
            for run in p.runs:
                if bold_all:
                    run.font.bold = True
                if color is not None:
                    run.font.color.rgb = color


def _repeat_header_row(row) -> None:
    """Повтор строки-шапки таблицы на каждой странице (длинные таблицы отчёта)."""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


_TABLE_TOTAL_WIDTH_CM = 26.7    # A4 ландшафт минус поля 15мм×2 (297мм-30мм), см. _setup_page
_TABLE_MIN_WIDTH_FRAC = 0.06    # пол доли ширины на колонку (узкие «№»/«Код» не схлопываются в 0)
_TABLE_MAX_WIDTH_FRAC = 0.40    # потолок доли — одна длинная «Формулировка» не съедает всю таблицу


def _column_widths(header: list[str], rows: list[list[str]]) -> list:
    """Автоширина колонок (порт `capabilities/export/xlsx.py::_autosize`, спека 08 §8): доля
    общей ширины таблицы пропорциональна максимальной длине содержимого колонки (заголовок и
    все строки; у ячеек с «<br>» — только первая строка переноса, иначе многострочные ячейки
    раздувают оценку непропорционально), с полом/потолком доли — см. `_TABLE_MIN/MAX_WIDTH_FRAC`."""
    ncols = len(header)
    if ncols == 0:
        return []
    lens = [max(1, len(str(header[i]))) for i in range(ncols)]
    for row in rows:
        for i in range(min(ncols, len(row))):
            first_line = str(row[i]).split("<br>")[0]
            lens[i] = max(lens[i], len(first_line))
    total_len = sum(lens)
    min_cm = _TABLE_TOTAL_WIDTH_CM * _TABLE_MIN_WIDTH_FRAC
    max_cm = _TABLE_TOTAL_WIDTH_CM * _TABLE_MAX_WIDTH_FRAC
    widths_cm = []
    for length in lens:
        raw = (_TABLE_TOTAL_WIDTH_CM * (length / total_len) if total_len
               else _TABLE_TOTAL_WIDTH_CM / ncols)
        widths_cm.append(min(max(raw, min_cm), max_cm))
    return [Cm(w) for w in widths_cm]


def _apply_column_widths(table, widths: list) -> None:
    """Проставить ширину каждой ячейке колонки (шапка + строки) — python-docx не держит ширину
    надёжно, если задать её только на `table.columns[i]` без каждой ячейки в отдельности."""
    if not widths:
        return
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]
    for i, column in enumerate(table.columns):
        if i < len(widths):
            column.width = widths[i]


def _add_table(doc: Document, header: list[str], rows: list[list[str]], palette: dict) -> None:
    """Таблица отчёта (типографика 06.07): шапка — жирная, акцентная, с фоном, повторяется на
    разрыве страниц; первая колонка — жирная на светлом фоне (навигационная); «зебра» строк;
    шрифт ячеек = основной − 1; в ячейках работают **жирный**, *курсив* и «<br>»; ширины колонок —
    автоматически по содержимому (`_column_widths`)."""
    base_size = 10.5 - _TABLE_FONT_DELTA
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        _fill_cell(hdr_cells[i], text, base_size, bold_all=True, color=_rgb(palette["accent"]))
        _shade_cell(hdr_cells[i], palette["block"])
    _repeat_header_row(table.rows[0])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row[: len(header)]):
            _fill_cell(cells[i], text, base_size, bold_all=(i == 0))
        # первая колонка — навигационная: всегда светлый фон; остальные — «зебра» нечётных строк
        _shade_cell(cells[0], "EDEDED")
        if row_idx % 2 == 1:
            for cell in cells[1:]:
                _shade_cell(cell, palette["block"])
    _apply_column_widths(table, _column_widths(header, rows))
    doc.add_paragraph()


def _shade_cell(cell, hex_color: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color.lstrip("#")})
    cell._tc.get_or_add_tcPr().append(shd)


def _add_image(doc: Document, caption: str, rel_path: str, workspace: Path | None) -> None:
    """Блок картинки «![подпись](путь)» (кадр демонстрации): путь — относительно воркспейса
    проекта; ширина ~160мм (в широкие поля A4-ландшафта помещается с запасом); подпись — курсивом
    под картинкой. Файл не найден — предупреждение текстом вместо падения всего документа."""
    full_path = Path(workspace) / rel_path if workspace else Path(rel_path)
    if full_path.is_file():
        doc.add_picture(str(full_path), width=Mm(160))
    else:
        doc.add_paragraph(f"[картинка не найдена: {rel_path}]")
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.italic = True
        run.font.size = Pt(9)
    doc.add_paragraph()


def _add_title_page(doc: Document, document_title: str, client: str | None, when: str | None,
                     palette: dict) -> None:
    """Титульный лист «корпоративного документа» (спека 05-b, редакция 2): название крупно,
    клиент/дата — если переданы (оба параметра опциональны), колонтитульная линия темы снизу."""
    for _ in range(4):
        doc.add_paragraph()
    # верхняя акцентная линия — рамка титула
    top_line = doc.add_paragraph()
    _add_bottom_border(top_line, palette["accent"])

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(document_title)
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = _rgb(palette["accent"])

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = sub_p.add_run(document_title)   # подзаголовок — «название» документа из конфига (спека 08 §8), не хардкод AS-IS
    sub.font.size = Pt(11)
    sub.font.color.rgb = _rgb(palette["text"])

    line_p = doc.add_paragraph()
    _add_bottom_border(line_p, palette["accent"])

    if client:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Подготовлено для: {client}")
        run.font.size = Pt(15)
        run.font.color.rgb = _rgb(palette["text"])

    if when:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(when)
        run.font.size = Pt(12)
        run.font.color.rgb = _rgb(palette["text"])

    for _ in range(8):
        doc.add_paragraph()
    conf_p = doc.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf = conf_p.add_run("КОНФИДЕНЦИАЛЬНО")
    conf.font.size = Pt(10)
    conf.font.bold = True
    conf.font.color.rgb = _rgb(palette["accent"])
    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = note_p.add_run("Документ содержит результаты обследования и предназначен только для адресата")
    note.font.size = Pt(8)
    note.font.italic = True
    note.font.color.rgb = _rgb(palette["text"])

    doc.add_page_break()


def _add_bottom_border(paragraph, hex_color: str) -> None:
    """Нижняя граница абзаца — «колонтитульная линия темы» на титульном листе."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color.lstrip("#"))
    borders.append(bottom)
    p_pr.append(borders)


def _add_toc_page(doc: Document, palette: dict) -> None:
    """Страница «Содержание» со стандартным полем Word TOC \\o "1-3" \\h \\z \\u.

    Поле TOC не считает содержимое сам python-docx — Word строит его при открытии документа
    (или по «Обновить поле») по своим стилям заголовков; предложение обновить при открытии —
    штатное поведение Word для документов с полями, не дефект рендера.
    """
    heading = doc.add_heading(level=1)
    heading.add_run("Содержание")
    _style_heading(heading, 1, palette)

    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Оглавление обновится при открытии документа (клик правой кнопкой → «Обновить поле»)."
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(placeholder)
    run._r.append(fld_char_end)

    doc.add_page_break()


def _add_paragraph_block(doc: Document, block: dict) -> None:
    """Абзац/маркированный пункт с поддержкой «**жирного**» (регулярный текст engine.py)."""
    if block["type"] == "bullet":
        p = doc.add_paragraph(style="List Bullet")
        _add_inline_runs(p, block["text"])
        return
    p = doc.add_paragraph()
    _add_inline_runs(p, block["text"])
    if block.get("italic"):
        for run in p.runs:
            run.font.italic = True


def _shade(paragraph, hex_color: str) -> None:
    """Заливка абзаца (фон врезки-callout)."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    p_pr.append(shd)


def _left_border(paragraph, hex_color: str) -> None:
    """Акцентная левая граница абзаца (маркер врезки-callout, приём Big4)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color.lstrip("#"))
    borders.append(left)
    p_pr.append(borders)


def _add_callout(doc: Document, lines: list[str], palette: dict) -> None:
    """Врезка-вывод (callout): заливка светло-серым + акцентная левая граница; первая строка —
    подзаголовком акцентного цвета. Приём отчётов Big4 — выносит ключевой вывод/эффект из потока."""
    for idx, text in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _shade(p, "F2F2F2")
        _left_border(p, palette["accent"])
        _add_inline_runs(p, text)
        if idx == 0:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = _rgb(palette["accent"])


def render_docx(md_text: str, out_path, project: str = "", document_title: str = "",
                 titul_client: str | None = None, titul_date: str | None = None,
                 титул: bool = True) -> Path:
    """Собрать docx-файл из канонического md. Тема — active_palette(project_path проекта).

    `titul_client`/`titul_date` — параметры титульного листа (спека 05-b, редакция 2), оба
    опциональны и влияют ТОЛЬКО на docx (канонический md, проверяемый на байт-стабильность,
    титула не несёт).

    `титул` (спека 08 §8, ключ конфига «документ.титул: нет») — рабочие/рассылочные документы
    (повестка, план встреч, открытые вопросы, гипотезы) НЕ несут титульный лист и «Содержание»:
    вызывающий код (cli.py/interface/documents.py) читает ключ из эффективного конфига документа
    и передаёт сюда `титул=False`; по умолчанию (без ключа/явного параметра) титул остаётся, как
    у корпоративных документов (протокол, отчёт AS-IS/TO-BE, протокол материалов).
    """
    out_path = Path(out_path)
    workspace = Path(project) if project and Path(str(project)).is_dir() else None
    palette = active_palette(str(workspace) if workspace else None)
    title = document_title or out_path.stem

    doc = Document()
    _setup_page(doc)
    _add_footer(doc, project or "OPP", title)

    text_color = _rgb(palette["text"])
    style = doc.styles["Normal"]
    style.font.name = _BASE_FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = text_color
    style.paragraph_format.space_after = Pt(6)

    if титул:
        _add_title_page(doc, title, titul_client, titul_date, palette)
        _add_toc_page(doc, palette)

    for block in _parse_blocks(md_text):
        if block["type"] == "heading":
            p = doc.add_heading(level=min(block["level"], 9))
            p.add_run(block["text"])
            _style_heading(p, block["level"], palette)
            if block["level"] == 2:
                _page_break_before(p)
        elif block["type"] == "table":
            _add_table(doc, block["header"], block["rows"], palette)
        elif block["type"] == "image":
            _add_image(doc, block["caption"], block["path"], workspace)
        elif block["type"] == "callout":
            _add_callout(doc, block["lines"], palette)
        else:
            _add_paragraph_block(doc, block)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
