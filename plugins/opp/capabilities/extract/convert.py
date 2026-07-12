"""Конвертеры форматов источников → markdown (бронза). Контент не теряем.

Лёгкие форматы — стандартная библиотека. `.xlsx/.xlsm` и `.docx` — python-native (openpyxl,
python-docx: ставятся в venv автоматически, внешних бинарей не требуют). Остальной офис/PDF/
изображения — системные инструменты (LibreOffice `soffice`, `pdftotext`, `tesseract`), если
установлены. Чего не умеем — честно помечаем `readable=false` с различимой причиной (нет
бинаря / не успел за таймаут / ошибка конвертации / битый файл), но узел всё равно создаётся
(ничего не теряем).
"""

from __future__ import annotations
import csv as _csv
import html as _html
import json
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from capabilities.extract.secrets import get_secret

LIGHT_TEXT = {".txt", ".log", ".md", ".markdown"}
XLSX_NATIVE = {".xlsx", ".xlsm"}  # openpyxl, без soffice
DOCX_NATIVE = {".docx"}  # python-docx, без soffice
OFFICE_TEXT = {".doc", ".odt", ".rtf", ".pptx", ".ppt"}  # fallback — soffice
OFFICE_SHEET = {".xls", ".xlsb", ".ods"}  # fallback — soffice
IMAGES = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tif", ".tiff", ".webp"}
MEDIA = {".wav", ".mp3", ".m4a", ".ogg", ".flac", ".mp4", ".mov", ".avi", ".mkv", ".webm"}
NEXT_STEP = {".mpp", ".mxl"}


def _have(tool: str) -> bool:
    return shutil.which(tool) is not None


def _read_text(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="replace")


def _cell(s: str) -> str:
    return s.replace("|", "\\|").replace("\n", " ").strip()


def _csv_to_md(p: Path, delim: str) -> str:
    rows = list(_csv.reader(_read_text(p).splitlines(), delimiter=delim))
    rows = [r for r in rows if r]
    if not rows:
        return ""
    head = rows[0]
    out = ["| " + " | ".join(_cell(c) for c in head) + " |",
           "|" + "|".join(["---"] * len(head)) + "|"]
    for r in rows[1:]:
        out.append("| " + " | ".join(_cell(c) for c in r) + " |")
    return "\n".join(out)


def _json_to_md(p: Path) -> str:
    try:
        data = json.loads(_read_text(p))
        return "```json\n" + json.dumps(data, ensure_ascii=False, indent=2) + "\n```"
    except Exception:
        return "```\n" + _read_text(p) + "\n```"


def _strip_html(text: str) -> str:
    text = re.sub(r"(?is)<(script|style).*?</\1>", "", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = _html.unescape(text)
    return re.sub(r"\n{3,}", "\n\n", re.sub(r"[ \t]{2,}", " ", text)).strip()


def _short_err(exc) -> str:
    """Кратко: текст ошибки, обрезанный для note реестра."""
    s = str(exc).strip()
    if not s:
        return "неизвестная ошибка"
    return (s[:157] + "…") if len(s) > 158 else s


def _soffice_timeout(p: Path) -> int:
    """180с + 15с/МБ, потолок 900с (спека 08 §7)."""
    size_mb = p.stat().st_size / (1024 * 1024)
    return min(900, int(180 + 15 * size_mb))


def _soffice(p: Path, to: str) -> tuple:
    """Конвертировать через LibreOffice. Возвращает (текст|None, note) — note различает причину
    отказа: нет бинаря / не успел за таймаут / ошибка конвертации (спека 08 §7)."""
    if not _have("soffice"):
        return None, "нужен LibreOffice (soffice) — см. УСТАНОВКА"
    timeout = _soffice_timeout(p)
    with tempfile.TemporaryDirectory() as d:
        try:
            subprocess.run(
                ["soffice", "--headless", f"-env:UserInstallation=file://{d}/prof",
                 "--convert-to", to, "--outdir", d, str(p)],
                check=True, capture_output=True, timeout=timeout)
        except subprocess.TimeoutExpired:
            size_mb = p.stat().st_size / (1024 * 1024)
            return None, (f"LibreOffice не успел (файл {size_mb:.1f} МБ, лимит {timeout} с) — "
                          "повторите или разбейте файл")
        except subprocess.CalledProcessError as exc:
            err = (exc.stderr or b"").decode("utf-8", "replace").strip() if exc.stderr else ""
            return None, f"ошибка конвертации LibreOffice: {_short_err(err) if err else 'см. код возврата soffice'}"
        except Exception as exc:  # noqa: BLE001
            return None, f"ошибка конвертации LibreOffice: {_short_err(exc)}"
        ext = to.split(":")[0]
        outs = list(Path(d).glob(f"*.{ext}"))
        if outs:
            return outs[0].read_text(encoding="utf-8", errors="replace"), ""
    return None, "ошибка конвертации LibreOffice: файл результата не создан"


def _pdf_text(p: Path):
    if not _have("pdftotext"):
        return None
    try:
        r = subprocess.run(["pdftotext", "-layout", str(p), "-"],
                           check=True, capture_output=True, timeout=180)
        return r.stdout.decode("utf-8", "replace")
    except Exception:
        return None


def _ocr(p: Path):
    if not _have("tesseract"):
        return None
    try:
        r = subprocess.run(["tesseract", str(p), "stdout", "-l", "rus+eng"],
                           check=True, capture_output=True, timeout=180)
        return r.stdout.decode("utf-8", "replace")
    except Exception:
        return None


def _pdf_ocr(p: Path):
    """OCR для PDF без текстового слоя (сканы): pdftoppm → tesseract постранично."""
    if not (_have("pdftoppm") and _have("tesseract")):
        return None
    with tempfile.TemporaryDirectory() as d:
        try:
            subprocess.run(["pdftoppm", "-r", "200", "-png", str(p), str(Path(d) / "стр")],
                           check=True, capture_output=True, timeout=600)
        except Exception:
            return None
        texts = []
        for img in sorted(Path(d).glob("стр-*.png")):
            try:
                r = subprocess.run(["tesseract", str(img), "stdout", "-l", "rus+eng"],
                                   check=True, capture_output=True, timeout=180)
                texts.append(r.stdout.decode("utf-8", "replace"))
            except Exception:
                pass
        joined = "\n\n".join(t for t in texts if t.strip())
        return joined or None


def _find_java(tool: str = "java"):
    exe = tool + (".exe" if os.name == "nt" else "")
    home = os.environ.get("JAVA_HOME")
    if home and (Path(home) / "bin" / exe).exists():
        return str(Path(home) / "bin" / exe)
    for c in ("/opt/homebrew/opt/openjdk@21", "/opt/homebrew/opt/openjdk@17",
              "/opt/homebrew/opt/openjdk", "/usr/local/opt/openjdk"):
        if (Path(c) / "bin" / exe).exists():
            return str(Path(c) / "bin" / exe)
    return shutil.which(tool)


def _mpxj_lib():
    """Каталог с jar-ами MPXJ (пакет mpxj несёт их) — БЕЗ import mpxj (он тянет jpype)."""
    import sysconfig
    for key in ("purelib", "platlib"):
        sp = sysconfig.get_paths().get(key)
        if sp:
            cand = Path(sp) / "mpxj" / "lib"
            if cand.is_dir() and any(cand.glob("*.jar")):
                return cand
    return None


def _mpp_to_md(p: Path):
    """MS Project → markdown через MPXJ (jar-ы пакета mpxj) + отдельный процесс java."""
    lib = _mpxj_lib()
    if lib is None:
        return None
    java = _find_java("java")
    if not java:
        return None
    here = Path(__file__).resolve().parent
    src, cls = here / "MppRead.java", here / "MppRead.class"
    jars = os.path.join(lib, "*")
    javac = _find_java("javac")
    if not cls.exists() and src.exists() and javac:
        subprocess.run([javac, "-cp", jars, "-d", str(here), str(src)],
                       capture_output=True, timeout=180)
    try:
        r = subprocess.run([java, "-Dstdout.encoding=UTF-8", "-cp", jars + os.pathsep + str(here),
                            "MppRead", str(p)], capture_output=True, text=True, timeout=300)
        return r.stdout.strip() or None
    except Exception:
        return None


def _direct_url(url: str) -> str:
    """Преобразовать ссылку облака в прямую ссылку на скачивание."""
    if "drive.google.com" in url:
        m = re.search(r"/d/([\w-]+)", url) or re.search(r"[?&]id=([\w-]+)", url)
        if m:
            return f"https://drive.google.com/uc?export=download&id={m.group(1)}"
    if "dropbox.com" in url:
        if "dl=" in url:
            return re.sub(r"dl=0", "dl=1", url)
        return url + ("&dl=1" if "?" in url else "?dl=1")
    if "disk.yandex" in url or "yadi.sk" in url:
        import urllib.parse
        import urllib.request
        api = ("https://cloud-api.yandex.net/v1/disk/public/resources/download?public_key="
               + urllib.parse.quote(url))
        try:
            with urllib.request.urlopen(api, timeout=30) as r:  # noqa: S310
                return json.loads(r.read())["href"]
        except Exception:
            return url
    return url


def download_url(url: str, dest_dir) -> Path:
    """Скачать файл по ссылке (с поддержкой облаков) в папку. Возвращает путь к файлу."""
    import urllib.parse
    import urllib.request
    direct = _direct_url(url)
    req = urllib.request.Request(direct, headers={"User-Agent": "OPP-ingest/1.0"})
    with urllib.request.urlopen(req, timeout=120) as r:  # noqa: S310
        cd = r.headers.get("Content-Disposition", "")
        name = None
        m = re.search(r"filename\*?=(?:UTF-8'')?\"?([^\";]+)", cd)
        if m:
            name = urllib.parse.unquote(m.group(1)).strip()
        if not name:
            name = os.path.basename(urllib.parse.urlparse(direct).path) or "download.bin"
        data = r.read(200_000_000)
    dest = Path(dest_dir)
    dest.mkdir(parents=True, exist_ok=True)
    out = dest / name
    out.write_bytes(data)
    return out


def _dbf_to_md(p: Path):
    import struct
    data = p.read_bytes()
    try:
        numrec = struct.unpack("<I", data[4:8])[0]
        hdrlen = struct.unpack("<H", data[8:10])[0]
        reclen = struct.unpack("<H", data[10:12])[0]
        fields, i = [], 32
        while i < len(data) and data[i] != 0x0D:
            name = data[i:i + 11].split(b"\x00")[0].decode("cp866", "replace")
            fields.append((name, data[i + 16]))
            i += 32
        out = ["| " + " | ".join(n for n, _ in fields) + " |",
               "|" + "|".join(["---"] * len(fields)) + "|"]
        for r in range(numrec):
            rec = data[hdrlen + r * reclen: hdrlen + (r + 1) * reclen]
            if not rec or rec[0:1] == b"*":
                continue
            pos, vals = 1, []
            for _name, flen in fields:
                vals.append(rec[pos:pos + flen].decode("cp866", "replace").strip().replace("|", "\\|"))
                pos += flen
            out.append("| " + " | ".join(vals) + " |")
        return "\n".join(out)
    except Exception:
        return None


def _drawio_to_md(p: Path):
    import base64
    import re
    import urllib.parse
    import zlib
    text = _read_text(p)
    xmls = [text]
    if "<mxGraphModel" not in text:
        for d in re.findall(r"<diagram[^>]*>([^<]+)</diagram>", text):
            try:
                dec = zlib.decompress(base64.b64decode(d), -15)
                xmls.append(urllib.parse.unquote(dec.decode("utf-8", "replace")))
            except Exception:
                pass
    labels = []
    for x in xmls:
        labels += re.findall(r'value="([^"]+)"', x)
    labels = list(dict.fromkeys(_html.unescape(l).strip() for l in labels if l.strip()))
    return ("## Метки диаграммы\n" + "\n".join(f"- {l}" for l in labels)) if labels else None


def _vsdx_to_md(p: Path):
    import re
    import zipfile
    try:
        texts = []
        with zipfile.ZipFile(p) as z:
            for n in z.namelist():
                if n.startswith("visio/pages/") and n.endswith(".xml"):
                    x = z.read(n).decode("utf-8", "replace")
                    for m in re.findall(r"<Text[^>]*>(.*?)</Text>", x, re.S):
                        t = _html.unescape(re.sub(r"<[^>]+>", "", m)).strip()
                        if t:
                            texts.append(t)
        return ("## Текст из Visio\n" + "\n".join(f"- {t}" for t in texts)) if texts else None
    except Exception:
        return None


def _xlsx_cell(v) -> str:
    return "" if v is None else str(v)


def _xlsx_native_to_md(p: Path) -> str:
    """.xlsx/.xlsm → markdown через openpyxl (без LibreOffice, спека 08 §7).

    По листу: «## <имя>» + md-таблица. Гигантские листы (данных строк > 300) — шапка +
    первые 50 строк + «… всего строк: N» (полный счёт — материал для описи полноты, §4).
    Битый файл — исключение наружу (ловит вызывающий, честный note без падения).
    """
    import openpyxl
    GIANT_ROWS, PREVIEW_ROWS = 300, 50
    wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    try:
        parts = []
        for name in wb.sheetnames:
            ws = wb[name]
            rows = [list(r) for r in ws.iter_rows(values_only=True)]
            while rows and all(c is None for c in rows[-1]):  # хвостовые пустые (форматирование)
                rows.pop()
            parts.append(f"## {name}")
            if not rows:
                parts.append("_(пустой лист)_")
                continue
            head, data = rows[0], rows[1:]
            total = len(data)
            truncated = total > GIANT_ROWS
            shown = data[:PREVIEW_ROWS] if truncated else data
            table = ["| " + " | ".join(_cell(_xlsx_cell(c)) for c in head) + " |",
                     "|" + "|".join(["---"] * len(head)) + "|"]
            for row in shown:
                table.append("| " + " | ".join(_cell(_xlsx_cell(c)) for c in row) + " |")
            parts.append("\n".join(table))
            if truncated:
                non_empty = sum(1 for row in data if any(c is not None for c in row))
                parts.append(f"\n… всего строк: {total} (непустых: {non_empty})")
        return "\n\n".join(parts)
    finally:
        wb.close()


def _docx_native_to_md(p: Path) -> str:
    """.docx → markdown через python-docx (без LibreOffice, спека 08 §7).

    Параграфы (заголовки по стилю Heading N/Title → #-уровни, грубо) + таблицы → md-таблицы.
    Битый файл — исключение наружу (ловит вызывающий, честный note без падения).
    """
    import docx
    document = docx.Document(str(p))
    parts = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name if para.style else "") or ""
        m = re.match(r"Heading (\d+)", style)
        if m:
            parts.append("#" * min(int(m.group(1)), 6) + " " + text)
        elif style == "Title":
            parts.append("# " + text)
        else:
            parts.append(text)
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        head, data = rows[0], rows[1:]
        md = ["| " + " | ".join(_cell(c) for c in head) + " |",
              "|" + "|".join(["---"] * len(head)) + "|"]
        for r in data:
            md.append("| " + " | ".join(_cell(c) for c in r) + " |")
        parts.append("\n".join(md))
    return "\n\n".join(parts)


def _xlsx_shapes_to_md(p: Path):
    """Тексты плавающих фигур (текстовые блоки, SmartArt) внутри .xlsx/.xlsm.

    LibreOffice при экспорте листа в HTML переносит только содержимое ячеек — плавающая
    графика физически теряется. DrawingML-разметка фигур лежит отдельно, в xl/drawings/*.xml.
    Только .xlsx/.xlsm — это OOXML/zip; .xls (OLE2), .xlsb (другой бинарный формат) и .ods
    (draw:frame вместо DrawingML) сюда не входят.
    """
    import re
    import zipfile
    try:
        texts = []
        with zipfile.ZipFile(p) as z:
            for n in z.namelist():
                if n.startswith("xl/drawings/drawing") and n.endswith(".xml"):
                    x = z.read(n).decode("utf-8", "replace")
                    # \b — иначе <a:t...> ложно матчит соседние теги с общим префиксом
                    # (a:tailEnd, a:tabLst и т.п.)
                    for m in re.findall(r"<a:t\b[^>]*>(.*?)</a:t>", x, re.S):
                        t = _html.unescape(re.sub(r"<[^>]+>", "", m)).strip()
                        if t:
                            texts.append(t)
        if not texts:
            return None
        return "## Текст графических объектов\n\n" + "\n".join(f"- {t}" for t in texts)
    except Exception:
        return None


def fetch_url(url: str) -> tuple[str, bool, str]:
    """Получить содержимое по ссылке → markdown (для SRC-WEB)."""
    import urllib.request
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "OPP-ingest/1.0"})
        with urllib.request.urlopen(req, timeout=30) as r:  # noqa: S310
            ct = r.headers.get("Content-Type", "")
            data = r.read(5_000_000)
        text = data.decode("utf-8", "replace")
        if "html" in ct.lower() or "<html" in text[:2000].lower():
            return _strip_html(text), True, ""
        return text, True, ""
    except Exception as exc:  # noqa: BLE001
        return "", False, f"не удалось получить ссылку: {exc}"


def to_markdown(path) -> tuple[str, bool, str]:
    """Вернуть (markdown, readable, note). readable=false → узел есть, но контент не извлечён."""
    p = Path(path)
    suf = p.suffix.lower()

    if suf in LIGHT_TEXT:
        return _read_text(p), True, ""
    if suf == ".csv":
        return _csv_to_md(p, ","), True, ""
    if suf == ".tsv":
        return _csv_to_md(p, "\t"), True, ""
    if suf == ".json":
        return _json_to_md(p), True, ""
    if suf == ".xml":
        return "```xml\n" + _read_text(p) + "\n```", True, ""
    if suf in (".html", ".htm"):
        return _strip_html(_read_text(p)), True, ""
    if suf in DOCX_NATIVE:
        try:
            return _docx_native_to_md(p), True, ""
        except Exception as exc:  # noqa: BLE001 — битый файл, не падение
            return "", False, f"не удалось прочитать ({_short_err(exc)})"
    if suf in XLSX_NATIVE:
        try:
            md = _xlsx_native_to_md(p)
        except Exception as exc:  # noqa: BLE001 — битый файл, не падение
            return "", False, f"не удалось прочитать ({_short_err(exc)})"
        shapes_md = _xlsx_shapes_to_md(p)
        if shapes_md:
            md = md + "\n\n" + shapes_md
        return md, True, ""
    if suf in OFFICE_TEXT:
        t, note = _soffice(p, "txt:Text")
        return (t, True, "") if t is not None else ("", False, note)
    if suf in OFFICE_SHEET:
        t, note = _soffice(p, "html:HTML")
        if t is None:
            return "", False, note
        return _strip_html(t), True, ""
    if suf == ".pdf":
        t = _pdf_text(p)
        if t is not None and t.strip():
            return t, True, ""
        ocr = _pdf_ocr(p)
        if ocr:
            return ocr, True, ""
        return "", False, "PDF без текстового слоя; нужны pdftotext + pdftoppm/tesseract (OCR)"
    if suf in IMAGES:
        t = _ocr(p)
        return ("```\n" + t + "\n```", True, "") if t is not None else ("", False, "нужен tesseract (OCR)")
    if suf == ".dbf":
        t = _dbf_to_md(p)
        return (t, True, "") if t else ("", False, "не удалось разобрать DBF")
    if suf == ".drawio":
        t = _drawio_to_md(p)
        return (t, True, "") if t else ("", False, "не удалось разобрать drawio")
    if suf == ".vsdx":
        t = _vsdx_to_md(p)
        return (t, True, "") if t else ("", False, "не удалось разобрать vsdx")
    if suf in MEDIA:
        return "", False, "запись встречи — обрабатывается как SRC-REC (транскрибация БИТ Ньютон)"
    if suf in (".mpp", ".mpx"):
        if not _mpp_ready():
            return "", False, "MS Project: нужны java (JDK 17+) и пакет mpxj (pip install mpxj)"
        t = _mpp_to_md(p)
        if t:
            return "## Задачи проекта (MS Project)\n\n" + t, True, ""
        return "", False, "не удалось разобрать файл MS Project (повреждён или неподдерживаемая версия)"
    if suf == ".mxl":
        return "", False, ("mxl — табличный формат 1С; конвертируется в xlsx/xml самой 1С "
                           "(набело-ридера нет)")
    # неизвестный формат — пробуем как текст
    try:
        return _read_text(p), True, ""
    except Exception:
        return "", False, f"неизвестный формат {suf}"


def adapter_status() -> list[dict]:
    """Статус адаптеров на этой машине — для окна."""
    newton_ok = _have("newton") and bool(get_secret("NEWTON_TOKEN"))
    return [
        {"группа": "Текст/таблицы/данные", "форматы": "txt, md, csv, tsv, json, xml, html",
         "статус": "работает"},
        {"группа": "Офис (native)", "форматы": "docx, xlsx, xlsm", "статус": "работает"},
        {"группа": "Офис (LibreOffice)", "форматы": "doc, xls, xlsb, pptx, ppt, ods, rtf",
         "статус": "работает" if _have("soffice") else "нужен LibreOffice"},
        {"группа": "PDF", "форматы": "pdf",
         "статус": "работает" if _have("pdftotext") else "нужен pdftotext"},
        {"группа": "Изображения (OCR)", "форматы": "jpg, png, tiff…",
         "статус": "работает" if _have("tesseract") else "нужен tesseract"},
        {"группа": "Схемы/устар. данные", "форматы": "vsdx, drawio, dbf",
         "статус": "работает"},
        {"группа": "Записи встреч", "форматы": "mp4, mov, mp3, wav…",
         "статус": "работает (Newton)" if newton_ok else "нужен NEWTON_TOKEN (Newton)"},
        {"группа": "MS Project", "форматы": "mpp, mpx",
         "статус": "работает (MPXJ)" if _mpp_ready() else "нужны java + mpxj"},
        {"группа": "1С табличный", "форматы": "mxl",
         "статус": "выгрузка из 1С (xlsx/xml)"},
        {"группа": "Ссылки/облака", "форматы": "сайты, Я.Диск, Google Drive, Dropbox",
         "статус": "работает (скачивание + разбор)"},
    ]


def _mpp_ready() -> bool:
    return _mpxj_lib() is not None and bool(_find_java("java"))
